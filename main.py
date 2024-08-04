from selenium import webdriver
from selenium.webdriver.common.by import By
from docx import Document
import os
import sqlite3

connection = sqlite3.connect('my_database.db')              #Подключение к бд
cursor = connection.cursor()
                                                            #Создаём таблицу в бд, если её ещё нет
cursor.execute('''CREATE TABLE IF NOT EXISTS News (         
id INT PRIMARY KEY, 
title TEXT NOT NULL, 
body TEXT NOT NULL)''')

doc = Document()                                            #Создём docx документ, куда будем записывать найденный текст
doc.add_heading('Украденные новости', level=1)

driver = webdriver.Firefox()                                #Подключаем драйвер для мазилы
driver.get("https://www.sports.ru/football/news/")          #Переходим на страницу по ссылке

elems = driver.find_elements(By.CLASS_NAME, "short-text")   #Поиск всех элементов на странице по имени класса short-text
links = []
for elem in elems:
    links.append(elem.get_attribute('href'))                      #Извлекаем из каждого найденного элемента ссылку

cursor.execute('SELECT max(id) FROM News')                      #Обеспечение уникальности идентификаторов в бд
max_id_el = cursor.fetchone()
max_id = max_id_el[0] + 1

for i in range(25):
    all_txt = ''
    driver.get(links[i])                                    #Поочерёдно переходим по ссылкам
    title_elem = driver.find_element(By.TAG_NAME, "h1")     #Берём элемент заголовока
    text_elems = driver.find_elements(By.TAG_NAME, "p")     #Берём элемент каждого абзаца
    doc.add_heading(title_elem.text, level=2)               #Записываем текст, найденный в элементе
    for j in range(3, len(text_elems)-1):                   #text_elems[0-4].text - пустые, text_elems[-1] - реклама
        all_txt += text_elems[j].text
        doc.add_paragraph(text_elems[j].text)
    doc.add_page_break()
    cursor.execute('INSERT INTO News (id, title, body) VALUES (?, ?, ?)', (max_id + i, title_elem.text, all_txt))

doc.save(os.path.dirname(os.path.abspath(__file__)))  #Сохраняем документ в файл

connection.commit()                                         #Сохраняем изменения в бд
connection.close()                                          #Закрываем подключение к бд
driver.close()
